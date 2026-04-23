from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy.orm import Session
from fastapi.responses import FileResponse, JSONResponse, Response

from ..deps import get_current_user
from ..db.session import get_db
from ..db.models import RepairTask, User, UserSatisfactionRating, OcrResult, CompletionLog, ModelUsageStat
from ..schemas.doc import (
    RepairResponse,
    UserSatisfactionRatingRequest,
    UserSatisfactionRatingResponse
)
from ..services.doc_repair import repair_document, fill_missing_text
from ..services.storage import save_result_image_bytes, save_upload_bytes, new_task_id
from ..services.pdf_processor import is_pdf_file, process_pdf

router = APIRouter(prefix="/api/doc", tags=["doc"])


# 测试用的简化认证依赖
def get_current_user_optional(current_user: Optional[User] = Depends(get_current_user)):
    """可选认证，用于测试"""
    return current_user


def _b64_to_bytes(b64: str) -> bytes:
    return base64.b64decode(b64)


def save_ocr_result(db: Session, task_id: int, user_id: int, ocr_result: dict, processing_time: float = 0.0):
    """
    保存 OCR 识别结果到数据库
    
    Args:
        db: 数据库会话
        task_id: 任务 ID
        user_id: 用户 ID
        ocr_result: OCR 识别结果字典
        processing_time: 处理时间（秒）
    """
    if not ocr_result or not ocr_result.get('success'):
        return
    
    text = ocr_result.get('text', '')
    confidences = ocr_result.get('confidences', [])
    
    # 计算统计信息
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    line_count = len(lines)
    word_count = len(text.replace('\n', '').replace(' ', ''))
    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
    
    # 创建 OCR 结果记录
    db_ocr_result = OcrResult(
        task_id=task_id,
        user_id=user_id,
        language='ch',  # 默认中文
        confidence=avg_confidence,
        word_count=word_count,
        line_count=line_count,
        raw_text=text,
        structured_data=ocr_result.get('structured_data', []),
        processing_time=processing_time,
        model_version='RapidOCR_v5'
    )
    
    db.add(db_ocr_result)
    # 注意：不在这里 commit，由调用者统一 commit


@router.post("/repair", response_model=RepairResponse)
async def repair(
    file: UploadFile = File(...),
    doc_type: str = Form("document"),
    use_super_resolution: str = Form("true"),
    sr_model_type: str = Form("classical"),
    sr_scale: int = Form(4),
    enable_text_optimization: str = Form("false"),
    enable_spell_correction: str = Form("true"),
    pdf_page: int = Form(1),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> RepairResponse:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Empty file")

    # 保存上传文件（用于调试/复算）
    upload_path = save_upload_bytes(contents, file.filename, user_id=current_user.id)

    # 处理内容
    processed_contents = contents
    processing_info = {}
    
    # 检查是否为 PDF 文件
    if is_pdf_file(file.filename):
        print(f"[PDF Processor] 检测到 PDF 文件，正在转换第 {pdf_page} 页...")
        try:
            processed_contents, processing_info = process_pdf(contents, file.filename, page=pdf_page)
            print(f"[PDF Processor] PDF 转换成功：{processing_info}")
        except Exception as e:
            print(f"[PDF Processor] PDF 转换失败：{e}")
            raise HTTPException(status_code=400, detail=f"PDF 处理失败：{str(e)}")
    
    # 检查是否为 ZIP 文件
    elif file.filename.lower().endswith('.zip'):
        return await handle_zip_file(contents, file.filename, doc_type, current_user, db, use_super_resolution, sr_model_type)
    
    # 处理图像文件
    task_key = new_task_id()
    try:
        use_super_resolution_flag = use_super_resolution.lower() == "true"
        enable_text_optimization_flag = enable_text_optimization.lower() == "true"
        # 图片修复模块默认跳过 OCR（图片修复不需要 OCR）
        enable_ocr_flag = False
        enable_spell_correction_flag = enable_spell_correction.lower() == "true"
        result = repair_document(processed_contents, filename=file.filename, doc_type=doc_type, use_super_resolution=use_super_resolution_flag, sr_model_type=sr_model_type, sr_scale=sr_scale, enable_text_optimization=enable_text_optimization_flag, enable_ocr=enable_ocr_flag, enable_spell_correction=enable_spell_correction_flag)

        # 保存原始图片（可能是裁剪或 PDF 转换后的）
        save_result_image_bytes(processed_contents, task_id=task_key, user_id=current_user.id, filename="original.png")
        
        # 保存修复后的图片
        repaired_b64 = result["repaired_image_base64"]
        repaired_bytes = _b64_to_bytes(repaired_b64)
        save_result_image_bytes(repaired_bytes, task_id=task_key, user_id=current_user.id, filename="repaired.png")

        # 合并处理信息到 meta 中
        meta = result.get("meta", {})
        if processing_info:
            meta["pdf_processing"] = processing_info

        task = RepairTask(
            task_key=task_key,
            user_id=current_user.id,
            original_filename=file.filename,
            status="done",
            ocr_text=result.get("ocr_text", "") or "",
            filled_text=result.get("filled_text", "") or "",
            original_relpath=f"{task_key}_original.png",  # 保存原图路径
            repaired_relpath=f"{task_key}_repaired.png",
            meta=meta,
        )
    except Exception:
        # 先保存失败状态
        task = RepairTask(
            task_key=task_key,
            user_id=current_user.id,
            original_filename=file.filename,
            status="failed",
            ocr_text="",
            filled_text="",
            repaired_relpath=None,
            meta={"error": "repair_failed", "pdf_processing": processing_info} if processing_info else {"error": "repair_failed"},
        )
        db.add(task)
        db.commit()
        raise

    # 用数据库自增 id 作为 task_key，先 flush 获取 id，然后 commit
    db.add(task)
    db.flush()  # 刷新到数据库，获取自增 ID
    task_id = task.id
    # 保持原始的 task_key，不要覆盖它
    # task_key = task.task_key
    
    # 记录超分模型使用情况
    if use_super_resolution_flag:
        from datetime import datetime
        # 打印 sr_model_type 的值，用于调试
        print(f"[Repair] sr_model_type: {sr_model_type}")
        
        # 标准化模型类型，确保大小写一致
        normalized_sr_model_type = sr_model_type.lower()
        print(f"[Repair] normalized_sr_model_type: {normalized_sr_model_type}")
        
        model_type_map = {
            "classical": "classical",
            "compressed": "compressed",
            "realworld": "realworld"
        }
        model_type = model_type_map.get(normalized_sr_model_type, "classical")
        print(f"[Repair] model_type: {model_type}")
        
        # 查找或创建模型使用记录
        model_stat = db.query(ModelUsageStat).filter(
            ModelUsageStat.user_id == current_user.id,
            ModelUsageStat.model_type == model_type
        ).first()
        
        if not model_stat:
            model_stat = ModelUsageStat(
                user_id=current_user.id,
                model_type=model_type,
                model_name=f"{sr_model_type}_sr",
                usage_count=0,
                success_count=0,
                total_processing_time=0.0
            )
            db.add(model_stat)
            db.flush()
            print(f"[Repair] 创建新的模型使用记录: {model_type}")
        else:
            print(f"[Repair] 找到现有模型使用记录: {model_type}, 当前计数: {model_stat.usage_count}")
        
        # 更新统计数据
        processing_time = meta.get("performance", {}).get("total_processing_time", 0.0)
        model_stat.usage_count += 1
        if task.status == "done":
            model_stat.success_count += 1
        model_stat.total_processing_time += processing_time
        print(f"[Repair] 更新模型使用记录: {model_type}, 新计数: {model_stat.usage_count}")
    
    # 保存 OCR 结果（使用实际的 task.id）
    ocr_raw_result = result.get("ocr_raw_result", {})
    if ocr_raw_result:
        save_ocr_result(db, task_id, current_user.id, ocr_raw_result, meta.get("performance", {}).get("total_processing_time", 0.0))
    
    # 在 commit 之前保存需要的字段值
    ocr_text = task.ocr_text
    filled_text = task.filled_text
    task_id_value = task.id
    task_key_value = task.task_key
    filename_value = task.original_filename
    status_value = task.status
    
    print(f"[Repair] 准备提交任务：ID={task_id_value}, task_key={task_key_value}, filename={filename_value}, status={status_value}")
    
    # 统一提交所有更改（包括 RepairTask 和 OcrResult）
    db.commit()
    
    print(f"[Repair] 任务提交成功：ID={task_id_value}, task_key={task_key_value}")
    
    # 返回响应：前端通常只需要 task_id 与结果
    response = RepairResponse(
        task_id=task_key_value,
        ocr_text=ocr_text,
        filled_text=filled_text,
        repaired_image_base64=result["repaired_image_base64"],
        mild_image_base64=result.get("mild_image_base64", result["repaired_image_base64"]),
        strong_image_base64=result.get("strong_image_base64", result["repaired_image_base64"]),
        repaired_image_mime="image/png",
        meta=meta,
    )
    
    # 添加缓存控制头，防止浏览器缓存旧的结果
    from fastapi.responses import JSONResponse
    return JSONResponse(
        content=response.dict(),
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
    )


async def handle_zip_file(
    zip_contents: bytes,
    zip_filename: str,
    doc_type: str,
    current_user: User,
    db: Session,
    use_super_resolution: str = "true",
    sr_model_type: str = "classical",
) -> JSONResponse:
    """处理 ZIP 文件，批量修复其中的图像"""
    try:
        # 解压 ZIP 文件
        with zipfile.ZipFile(io.BytesIO(zip_contents), 'r') as zip_ref:
            # 检查 ZIP 文件中的文件
            zip_info_list = zip_ref.infolist()
            image_files = []
            
            for info in zip_info_list:
                # 跳过目录
                if info.filename.endswith('/'):
                    continue
                # 只处理图像文件
                if any(info.filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']):
                    image_files.append(info)
            
            if not image_files:
                raise HTTPException(status_code=400, detail="No image files found in ZIP")
            
            # 处理每个图像文件
            results = []
            for info in image_files:
                with zip_ref.open(info) as file:
                    image_contents = file.read()
                    task_key = new_task_id()
                    
                    try:
                        use_super_resolution_flag = use_super_resolution.lower() == "true"
                        result = repair_document(image_contents, filename=info.filename, doc_type=doc_type, use_super_resolution=use_super_resolution_flag, sr_model_type=sr_model_type)
                        
                        # 保存原始图片
                        save_result_image_bytes(image_contents, task_id=task_key, user_id=current_user.id, filename=f"original_{info.filename}")
                        
                        # 保存修复后的图片
                        repaired_b64 = result["repaired_image_base64"]
                        repaired_bytes = _b64_to_bytes(repaired_b64)
                        save_result_image_bytes(repaired_bytes, task_id=task_key, user_id=current_user.id, filename=f"repaired_{info.filename}")
                        
                        # 合并处理信息到 meta 中
                        meta = result.get("meta", {})
                        
                        task = RepairTask(
                            task_key=task_key,
                            user_id=current_user.id,
                            original_filename=info.filename,
                            status="done",
                            ocr_text=result.get("ocr_text", "") or "",
                            filled_text=result.get("filled_text", "") or "",
                            original_relpath=f"{task_key}_original_{info.filename}",  # 保存原图路径
                            repaired_relpath=f"{task_key}_repaired_{info.filename}",
                            meta=meta,
                        )
                    except Exception:
                        # 保存失败状态
                        task = RepairTask(
                            task_key=task_key,
                            user_id=current_user.id,
                            original_filename=info.filename,
                            status="failed",
                            ocr_text="",
                            filled_text="",
                            repaired_relpath=None,
                            meta={"error": "repair_failed"},
                        )
                    
                    # 用数据库自增 id 作为 task_key，先 flush 获取 id，然后 commit
                    db.add(task)
                    db.flush()  # 刷新到数据库，获取自增 ID
                    task_id_value = task.id
                    
                    # 记录超分模型使用情况
                    if use_super_resolution_flag:
                        from datetime import datetime
                        # 打印 sr_model_type 的值，用于调试
                        print(f"[HandleZip] sr_model_type: {sr_model_type}")
                        
                        # 标准化模型类型，确保大小写一致
                        normalized_sr_model_type = sr_model_type.lower()
                        print(f"[HandleZip] normalized_sr_model_type: {normalized_sr_model_type}")
                        
                        model_type_map = {
                            "classical": "classical",
                            "compressed": "compressed",
                            "realworld": "realworld"
                        }
                        model_type = model_type_map.get(normalized_sr_model_type, "classical")
                        print(f"[HandleZip] model_type: {model_type}")
                        
                        # 查找或创建模型使用记录
                        model_stat = db.query(ModelUsageStat).filter(
                            ModelUsageStat.user_id == current_user.id,
                            ModelUsageStat.model_type == model_type
                        ).first()
                        
                        if not model_stat:
                            model_stat = ModelUsageStat(
                                user_id=current_user.id,
                                model_type=model_type,
                                model_name=f"{sr_model_type}_sr",
                                usage_count=0,
                                success_count=0,
                                total_processing_time=0.0
                            )
                            db.add(model_stat)
                            db.flush()
                            print(f"[HandleZip] 创建新的模型使用记录: {model_type}")
                        else:
                            print(f"[HandleZip] 找到现有模型使用记录: {model_type}, 当前计数: {model_stat.usage_count}")
                        
                        # 更新统计数据
                        processing_time = meta.get("performance", {}).get("total_processing_time", 0.0)
                        model_stat.usage_count += 1
                        if task.status == "done":
                            model_stat.success_count += 1
                        model_stat.total_processing_time += processing_time
                        print(f"[HandleZip] 更新模型使用记录: {model_type}, 新计数: {model_stat.usage_count}")
                    
                    db.commit()
                    db.refresh(task)
                    
                    results.append({
                        "task_id": task.task_key,
                        "filename": info.filename,
                        "status": task.status,
                        "ocr_text": task.ocr_text,
                        "filled_text": task.filled_text,
                    })
            
            # 返回批量处理结果
            return JSONResponse(
                status_code=200,
                content={
                    "message": f"Batch processing completed: {len(results)} images processed",
                    "results": results
                }
            )
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Invalid ZIP file")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing ZIP file: {str(e)}")


@router.get("/result/{task_id}")
def get_result_image(
    task_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    task = db.query(RepairTask).filter(RepairTask.task_key == task_id).first()
    if not task or task.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Task not found")
    if not task.repaired_relpath:
        raise HTTPException(status_code=404, detail="Result image not found")

    # repaired_relpath 目前仅用于占位：我们实际从 storage/results 找最后保存的文件
    from ..core.settings import storage_root

    root = storage_root() / "results"
    # 先在 user_id 子目录中查找
    user_results_dir = root / str(current_user.id)
    if user_results_dir.exists():
        for p in user_results_dir.glob(f"{task_id}_*.png"):
            return FileResponse(str(p))
    # 如果找不到，再在根目录查找（兼容旧版本）
    for p in root.glob(f"{task_id}_*.png"):
        return FileResponse(str(p))
    raise HTTPException(status_code=404, detail="Result image file missing")


@router.post("/repair-test")
async def repair_test(
    file: UploadFile = File(...),
    doc_type: str = Form("document"),
    use_super_resolution: str = Form("true"),
    sr_model_type: str = Form("classical"),
    pdf_page: int = Form(1),
    db: Session = Depends(get_db),
) -> RepairResponse:
    # 打印所有参数，用于调试
    print(f"[RepairTest] 接收到的参数:")
    print(f"[RepairTest] file.filename: {file.filename}")
    print(f"[RepairTest] doc_type: {doc_type}")
    print(f"[RepairTest] use_super_resolution: {use_super_resolution}")
    print(f"[RepairTest] sr_model_type: {sr_model_type}")
    print(f"[RepairTest] pdf_page: {pdf_page}")
    """
    测试用的修复接口，无需登录认证
    仅用于开发和测试阶段
    任务会保存到数据库，关联到测试用户（ID=1）
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    # 打印文件名，用于调试
    print(f"[RepairTest] 接收到的文件名: {file.filename}")
    print(f"[RepairTest] 文件名是否以 .pdf 结尾: {file.filename.lower().endswith('.pdf')}")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Empty file")

    # 处理内容
    processed_contents = contents
    processing_info = {}
    
    # 检查是否为PDF文件
    is_pdf = is_pdf_file(file.filename)
    print(f"[RepairTest] is_pdf_file 函数返回值: {is_pdf}")
    
    if is_pdf:
        print(f"[PDF Processor] 检测到PDF文件，正在转换第 {pdf_page} 页...")
        try:
            processed_contents, processing_info = process_pdf(contents, file.filename, page=pdf_page)
            print(f"[PDF Processor] PDF转换成功: {processing_info}")
        except Exception as e:
            print(f"[PDF Processor] PDF转换失败: {e}")
            raise HTTPException(status_code=400, detail=f"PDF处理失败: {str(e)}")

    use_super_resolution_flag = use_super_resolution.lower() == "true"
    
    # 打印 sr_model_type 的值，用于调试
    print(f"[RepairTest] 接收到的 sr_model_type: {sr_model_type}")
    
    # 获取或创建测试用户（ID=1）
    test_user = db.query(User).filter(User.id == 1).first()
    if not test_user:
        # 如果没有ID=1的用户，使用第一个用户
        test_user = db.query(User).first()
    
    if not test_user:
        raise HTTPException(status_code=500, detail="No user found in database")
    
    task_key = new_task_id()
    
    try:
        result = repair_document(
            processed_contents, 
            filename=file.filename, 
            doc_type=doc_type, 
            use_super_resolution=use_super_resolution_flag,
            sr_model_type=sr_model_type
        )
        
        # 保存修复后的图像
        repaired_b64 = result["repaired_image_base64"]
        repaired_bytes = _b64_to_bytes(repaired_b64)
        save_result_image_bytes(repaired_bytes, task_id=task_key, user_id=test_user.id, filename="repaired.png")
        
        # 合并处理信息到 meta 中
        meta = result.get("meta", {})
        meta["use_super_resolution"] = use_super_resolution_flag
        meta["sr_model_type"] = sr_model_type  # 保存使用的模型类型
        if processing_info:
            meta["pdf_processing"] = processing_info
        
        # 创建任务记录
        task = RepairTask(
            task_key=task_key,
            user_id=test_user.id,
            original_filename=file.filename,
            status="done",
            ocr_text=result.get("ocr_text", "") or "",
            filled_text=result.get("filled_text", "") or "",
            repaired_relpath=f"{task_key}_repaired.png",
            meta=meta,
        )
        
        # 记录超分模型使用情况
        if use_super_resolution_flag:
            from datetime import datetime
            # 打印 sr_model_type 的值，用于调试
            print(f"[RepairTest] sr_model_type: {sr_model_type}")
            
            # 标准化模型类型，确保大小写一致
            normalized_sr_model_type = sr_model_type.lower()
            print(f"[RepairTest] normalized_sr_model_type: {normalized_sr_model_type}")
            
            model_type_map = {
                "classical": "classical",
                "compressed": "compressed",
                "realworld": "realworld"
            }
            model_type = model_type_map.get(normalized_sr_model_type, "classical")
            print(f"[RepairTest] model_type: {model_type}")
            
            # 查找或创建模型使用记录
            model_stat = db.query(ModelUsageStat).filter(
                ModelUsageStat.user_id == test_user.id,
                ModelUsageStat.model_type == model_type
            ).first()
            
            if not model_stat:
                model_stat = ModelUsageStat(
                    user_id=test_user.id,
                    model_type=model_type,
                    model_name=f"{sr_model_type}_sr",
                    usage_count=0,
                    success_count=0,
                    total_processing_time=0.0
                )
                db.add(model_stat)
                print(f"[RepairTest] 创建新的模型使用记录: {model_type}")
            else:
                print(f"[RepairTest] 找到现有模型使用记录: {model_type}, 当前计数: {model_stat.usage_count}")
            
            # 更新统计数据
            processing_time = meta.get("performance", {}).get("total_processing_time", 0.0)
            model_stat.usage_count += 1
            if task.status == "done":
                model_stat.success_count += 1
            model_stat.total_processing_time += processing_time
            print(f"[RepairTest] 更新模型使用记录: {model_type}, 新计数: {model_stat.usage_count}")
        
        # 保存任务记录
        db.add(task)
        
        # 在 commit 之前保存需要的属性值
        task_key_value = task.task_key
        ocr_text_value = result.get("ocr_text", "")
        filled_text_value = result.get("filled_text", "")
        repaired_image_b64_value = result["repaired_image_base64"]
        mild_image_b64_value = result.get("mild_image_base64", result["repaired_image_base64"])
        strong_image_b64_value = result.get("strong_image_base64", result["repaired_image_base64"])
        
        db.commit()

        response = RepairResponse(
            task_id=task_key_value,
            ocr_text=ocr_text_value,
            filled_text=filled_text_value,
            repaired_image_base64=repaired_image_b64_value,
            mild_image_base64=mild_image_b64_value,
            strong_image_base64=strong_image_b64_value,
            repaired_image_mime="image/png",
            meta=meta,
        )
        
        # 添加缓存控制头，防止浏览器缓存旧的结果
        return JSONResponse(
            content=response.dict(),
            headers={
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0"
            }
        )
    except Exception as e:
        # 保存失败状态
        meta = {"error": str(e)}
        if processing_info:
            meta["pdf_processing"] = processing_info
        
        task = RepairTask(
            task_key=task_key,
            user_id=test_user.id,
            original_filename=file.filename,
            status="failed",
            ocr_text="",
            filled_text="",
            repaired_relpath=None,
            meta=meta,
        )
        db.add(task)
        db.commit()
        raise HTTPException(status_code=500, detail=f"Repair failed: {str(e)}")


@router.get("/export/word/{task_id}")
def export_word_report(
    task_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """导出为 Word 报告"""
    task = db.query(RepairTask).filter(RepairTask.task_key == task_id).first()
    if not task or task.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Task not found")
    
    try:
        from docx import Document
        from docx.shared import Inches
        import io
        
        # 创建 Word 文档
        doc = Document()
        doc.add_heading('老旧文档修复报告', 0)
        
        # 添加基本信息
        doc.add_heading('基本信息', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.cell(0, 0).text = '任务 ID'
        table.cell(0, 1).text = task.task_key
        table.cell(1, 0).text = '原始文件名'
        table.cell(1, 1).text = task.original_filename
        table.cell(2, 0).text = '文档类型'
        table.cell(2, 1).text = task.doc_type
        table.cell(3, 0).text = '处理状态'
        table.cell(3, 1).text = task.status
        
        # 添加 OCR 识别结果
        doc.add_heading('OCR 识别结果', level=1)
        doc.add_paragraph(task.ocr_text or '无识别结果')
        
        # 添加缺字补全结果
        doc.add_heading('缺字补全结果', level=1)
        doc.add_paragraph(task.filled_text or '无补全结果')
        
        # 保存为字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # 返回文件响应
        from fastapi.responses import Response
        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename=repair_report_{task_id}.docx'
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating Word report: {str(e)}")


@router.post("/refill")
async def refill_text(
    ocr_text: str = Form(...),
    current_user: User = Depends(get_current_user),
):
    """
    基于用户校对后的OCR文本，重新进行缺字补全
    """
    try:
        filled_text = fill_missing_text(ocr_text)
        return {
            "success": True,
            "filled_text": filled_text
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Refill failed: {str(e)}")


@router.post("/rating", response_model=UserSatisfactionRatingResponse)
async def submit_rating(
    rating: UserSatisfactionRatingRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """提交用户满意度评价"""
    # 查找任务
    task = db.query(RepairTask).filter(RepairTask.task_key == rating.task_id).first()
    if not task or task.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Task not found")
    
    # 检查是否已经评过分
    existing_rating = db.query(UserSatisfactionRating).filter(
        UserSatisfactionRating.task_id == task.id,
        UserSatisfactionRating.user_id == current_user.id
    ).first()
    
    if existing_rating:
        # 更新现有评分
        existing_rating.visual_quality = rating.visual_quality
        existing_rating.text_readability = rating.text_readability
        existing_rating.stain_removal = rating.stain_removal
        existing_rating.color_restoration = rating.color_restoration
        existing_rating.overall_satisfaction = rating.overall_satisfaction
        existing_rating.comments = rating.comments
    else:
        # 创建新评分
        new_rating = UserSatisfactionRating(
            task_id=task.id,
            user_id=current_user.id,
            visual_quality=rating.visual_quality,
            text_readability=rating.text_readability,
            stain_removal=rating.stain_removal,
            color_restoration=rating.color_restoration,
            overall_satisfaction=rating.overall_satisfaction,
            comments=rating.comments
        )
        db.add(new_rating)
        existing_rating = new_rating
    
    db.commit()
    db.refresh(existing_rating)
    
    return UserSatisfactionRatingResponse(
        id=existing_rating.id,
        task_id=task.task_key,
        visual_quality=existing_rating.visual_quality,
        text_readability=existing_rating.text_readability,
        stain_removal=existing_rating.stain_removal,
        color_restoration=existing_rating.color_restoration,
        overall_satisfaction=existing_rating.overall_satisfaction,
        comments=existing_rating.comments,
        created_at=existing_rating.created_at.isoformat() if existing_rating.created_at else ""
    )


@router.get("/rating/{task_id}", response_model=List[UserSatisfactionRatingResponse])
def get_task_ratings(
    task_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取指定任务的所有用户评价"""
    # 查找任务
    task = db.query(RepairTask).filter(RepairTask.task_key == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    
    # 获取该任务的所有评价
    ratings = db.query(UserSatisfactionRating).filter(
        UserSatisfactionRating.task_id == task.id
    ).order_by(UserSatisfactionRating.created_at.desc()).all()
    
    # 构建响应
    result = []
    for r in ratings:
        result.append(UserSatisfactionRatingResponse(
            id=r.id,
            task_id=task.task_key,
            user_id=r.user_id,
            visual_quality=r.visual_quality,
            text_readability=r.text_readability,
            stain_removal=r.stain_removal,
            color_restoration=r.color_restoration,
            overall_satisfaction=r.overall_satisfaction,
            comments=r.comments,
            created_at=r.created_at.isoformat() if r.created_at else ""
        ))
    
    return result


@router.get("/rating/stats/summary", response_model=dict)
def get_rating_summary(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取用户满意度统计（用于仪表盘）"""
    from sqlalchemy import func
    
    # 统计总体数据
    stats = db.query(
        func.count(UserSatisfactionRating.id).label('total_ratings'),
        func.avg(UserSatisfactionRating.visual_quality).label('avg_visual'),
        func.avg(UserSatisfactionRating.text_readability).label('avg_text'),
        func.avg(UserSatisfactionRating.stain_removal).label('avg_stain'),
        func.avg(UserSatisfactionRating.color_restoration).label('avg_color'),
        func.avg(UserSatisfactionRating.overall_satisfaction).label('avg_overall')
    ).first()
    
    return {
        'total_ratings': stats.total_ratings or 0,
        'avg_visual_quality': round(stats.avg_visual or 0, 2),
        'avg_text_readability': round(stats.avg_text or 0, 2),
        'avg_stain_removal': round(stats.avg_stain or 0, 2),
        'avg_color_restoration': round(stats.avg_color or 0, 2),
        'avg_overall_satisfaction': round(stats.avg_overall or 0, 2)
    }



@router.get("/rating/{task_id}")
async def get_rating(
    task_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取任务的用户满意度评价"""
    task = db.query(RepairTask).filter(RepairTask.task_key == task_id).first()
    if not task or task.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Task not found")
    
    rating = db.query(UserSatisfactionRating).filter(
        UserSatisfactionRating.task_id == task.id,
        UserSatisfactionRating.user_id == current_user.id
    ).first()
    
    if not rating:
        return {"has_rated": False}
    
    return {
        "has_rated": True,
        "rating": UserSatisfactionRatingResponse(
            id=rating.id,
            task_id=task.task_key,
            visual_quality=rating.visual_quality,
            text_readability=rating.text_readability,
            stain_removal=rating.stain_removal,
            color_restoration=rating.color_restoration,
            overall_satisfaction=rating.overall_satisfaction,
            comments=rating.comments,
            created_at=rating.created_at.isoformat() if rating.created_at else ""
        )
    }


@router.get("/rating-stats")
async def get_rating_stats(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取用户的评分统计"""
    ratings = db.query(UserSatisfactionRating).filter(
        UserSatisfactionRating.user_id == current_user.id
    ).all()
    
    if not ratings:
        return {"total_ratings": 0, "average_scores": {}}
    
    total = len(ratings)
    avg_visual = sum(r.visual_quality for r in ratings) / total
    avg_readability = sum(r.text_readability for r in ratings) / total
    avg_stain = sum(r.stain_removal for r in ratings) / total
    avg_color = sum(r.color_restoration for r in ratings) / total
    avg_overall = sum(r.overall_satisfaction for r in ratings) / total
    
    return {
        "total_ratings": total,
        "average_scores": {
            "visual_quality": round(avg_visual, 2),
            "text_readability": round(avg_readability, 2),
            "stain_removal": round(avg_stain, 2),
            "color_restoration": round(avg_color, 2),
            "overall_satisfaction": round(avg_overall, 2)
        }
    }




